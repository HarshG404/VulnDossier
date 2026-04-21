"""
docx_builder.py — Word report with:
  • Company branding (logo, colors, company name from config)
  • Table of Contents page 2 with clickable bookmarks
  • Each vulnerability on its own page (page break before each)
  • All 8 CVSS v3.1 metrics (incl. Integrity & Availability)
  • POC placeholder under every finding
"""
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from core.cvss import cvss_breakdown, sev_text_rgb, sev_bg_rgb, sev_badge_rgb
from config.config_manager import load_config, hex_to_rgb


# ── Color utilities ───────────────────────────────────────────────────────────

def _hex_to_rgb_color(h: str) -> RGBColor:
    try:
        r, g, b = hex_to_rgb(h)
        return RGBColor(r, g, b)
    except Exception:
        return RGBColor(29, 158, 117)


def _sev_rgb(sev):   return RGBColor(*sev_text_rgb(sev))
def _sev_bg(sev):    return RGBColor(*sev_bg_rgb(sev))
def _sev_badge(sev): return RGBColor(*sev_badge_rgb(sev))


# ── XML helpers ───────────────────────────────────────────────────────────────

def _set_cell_bg(cell, rgb: RGBColor):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}")
    tcPr.append(shd)


def _add_bookmark(paragraph, bookmark_id: int, name: str):
    """Insert a named bookmark at the start of a paragraph."""
    p = paragraph._p
    bm_start = OxmlElement("w:bookmarkStart")
    bm_start.set(qn("w:id"),   str(bookmark_id))
    bm_start.set(qn("w:name"), name)
    bm_end = OxmlElement("w:bookmarkEnd")
    bm_end.set(qn("w:id"), str(bookmark_id))
    p.insert(0, bm_start)
    p.append(bm_end)


def _add_hyperlink_to_bookmark(paragraph, text: str, bookmark_name: str,
                               font_size=10, bold=False):
    """Add a clickable internal hyperlink pointing to a named bookmark."""
    run_elem = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), "58A6FF")
    rPr.append(color_el)

    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(font_size * 2)))
    rPr.append(sz)

    if bold:
        b_el = OxmlElement("w:b")
        rPr.append(b_el)

    u_el = OxmlElement("w:u")
    u_el.set(qn("w:val"), "single")
    rPr.append(u_el)

    run_elem.append(rPr)
    t_el = OxmlElement("w:t")
    t_el.text = text
    t_el.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    run_elem.append(t_el)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), bookmark_name)
    hyperlink.append(run_elem)
    paragraph._p.append(hyperlink)


def _add_page_break(doc):
    """Add a manual page break paragraph."""
    p = doc.add_paragraph()
    run = p.add_run()
    br  = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


def _add_rule(doc, color_hex="1D9E75"):
    p   = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(8)
    pr  = p._p
    pPr = pr.get_or_add_pPr()
    pb  = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color_hex.lstrip("#"))
    pb.append(bot)
    pPr.append(pb)


def _para_shade(paragraph, rgb: RGBColor):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}")
    pPr.append(shd)


# ── Severity helpers ──────────────────────────────────────────────────────────

def _sev_counts(findings):
    c = {"Critical":0,"High":0,"Medium":0,"Low":0,"Informational":0}
    for f in findings:
        s = f.get("severity","Low")
        if s in c: c[s] += 1
    return c


def _overall_risk(findings):
    c = _sev_counts(findings)
    for s in ["Critical","High","Medium","Low","Informational"]:
        if c[s] > 0: return s
    return "Informational"


# ── Main builder ──────────────────────────────────────────────────────────────

def build_docx(report_data: dict, output_path: str) -> str:
    cfg      = load_config()
    meta     = report_data.get("meta", {})
    findings = report_data.get("findings", [])
    exec_sum = report_data.get("executive_summary", "")
    struct   = cfg.get("report_structure", {})
    company  = cfg.get("company", {})
    cover_c  = cfg.get("cover_page", {})
    branding = cfg.get("branding", {})

    accent_hex  = cfg.get("colors",{}).get("accent",   "#1D9E75").lstrip("#")
    hdr_bg_hex  = cfg.get("colors",{}).get("header_bg","#0D1117").lstrip("#")
    ACCENT      = _hex_to_rgb_color("#" + accent_hex)
    HDR_BG      = _hex_to_rgb_color("#" + hdr_bg_hex)
    DARK        = RGBColor(13,17,23)
    WHITE       = RGBColor(255,255,255)
    MUTED       = RGBColor(100,110,120)
    LIGHT_GRAY  = RGBColor(246,248,250)
    BORDER_RGB  = RGBColor(208,215,222)
    BODY_TEXT   = _hex_to_rgb_color(cfg.get("colors",{}).get("body_text","#1F2328"))

    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)

    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Running header/footer ────────────────────────────────────────────────
    hf_cfg = cfg.get("header_footer", {})
    if hf_cfg.get("show_header", True):
        header = doc.sections[0].header
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.clear()
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _sub_tmpl = lambda t: (t or "") \
            .replace("{client_name}",     meta.get("client_name","")) \
            .replace("{assessment_type}", meta.get("project_type","")) \
            .replace("{company_name}",    company.get("name","")) \
            .replace("{date}",            datetime.today().strftime("%Y-%m-%d"))

        hl = _sub_tmpl(hf_cfg.get("header_left", "{client_name}"))
        hr = _sub_tmpl(hf_cfg.get("header_right", "{assessment_type}"))
        r  = hp.add_run(f"{hl}   |   {hr}")
        r.font.size  = Pt(8)
        r.font.color.rgb = MUTED

    if hf_cfg.get("show_footer", True):
        footer = doc.sections[0].footer
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.clear()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fl = _sub_tmpl(hf_cfg.get("footer_left",  "CONFIDENTIAL — {client_name}"))
        fc = _sub_tmpl(hf_cfg.get("footer_center", "{company_name}"))
        r2 = fp.add_run(f"{fl}   |   {fc}")
        r2.font.size = Pt(8)
        r2.font.color.rgb = MUTED
        if hf_cfg.get("show_page_numbers", True):
            fp.add_run("   |   Page ")
            fldChar1 = OxmlElement("w:fldChar")
            fldChar1.set(qn("w:fldCharType"), "begin")
            instrText = OxmlElement("w:instrText")
            instrText.text = "PAGE"
            fldChar2 = OxmlElement("w:fldChar")
            fldChar2.set(qn("w:fldCharType"), "separate")
            fldChar3 = OxmlElement("w:fldChar")
            fldChar3.set(qn("w:fldCharType"), "end")
            run_el = fp.add_run()._r
            for el in [fldChar1, instrText, fldChar2, fldChar3]:
                run_el.append(el)

    bm_counter = [0]

    def next_bm():
        bm_counter[0] += 1
        return bm_counter[0]

    # ══════════════════════════════════════════════════════════════════
    # COVER PAGE
    # ══════════════════════════════════════════════════════════════════
    # Logo
    logo_path = branding.get("logo_path","")
    if cover_c.get("show_logo",True) and logo_path and os.path.exists(logo_path):
        try:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            lw  = Cm(branding.get("logo_width_mm",40)/10)
            lh  = Cm(branding.get("logo_height_mm",20)/10)
            run.add_picture(logo_path, width=lw, height=lh)
        except Exception:
            pass

    # Company name
    if cover_c.get("show_company_name",True) and company.get("name"):
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cp.add_run(company["name"])
        cr.bold = True
        cr.font.size = Pt(16)
        cr.font.color.rgb = ACCENT

    if cover_c.get("show_company_tagline",True) and company.get("tagline"):
        tp = doc.add_paragraph()
        tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tr = tp.add_run(company["tagline"])
        tr.font.size = Pt(10)
        tr.font.color.rgb = MUTED

    doc.add_paragraph()

    # Classification
    if cover_c.get("classification_banner",True):
        clf = doc.add_paragraph()
        clf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr2 = clf.add_run(f"[ {meta.get('classification','CONFIDENTIAL')} ]")
        cr2.bold = True
        cr2.font.size = Pt(11)
        cr2.font.color.rgb = RGBColor(180,0,0)

    doc.add_paragraph()

    # Assessment type title
    title_text = (cover_c.get("custom_cover_title") or
                  meta.get("project_type","Penetration Test Report")).upper()
    tp2 = doc.add_paragraph()
    tp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr2 = tp2.add_run(title_text)
    tr2.bold = True
    tr2.font.size = Pt(13)
    tr2.font.color.rgb = ACCENT

    # Client name
    clp = doc.add_paragraph()
    clp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    clr = clp.add_run(meta.get("client_name","Client Report"))
    clr.bold = True
    clr.font.size = Pt(24)
    clr.font.color.rgb = DARK

    doc.add_paragraph()

    # Meta table
    mt_rows = [
        ("Assessor",         meta.get("pentester_name","")),
        ("Client",           meta.get("client_name","")),
        ("Assessment Type",  meta.get("project_type","")),
        ("Scope",            meta.get("scope","")),
        ("Period",           f"{meta.get('start_date','')} → {meta.get('end_date','')}"),
        ("Classification",   meta.get("classification","CONFIDENTIAL")),
        ("Report Date",      datetime.today().strftime("%B %d, %Y")),
    ]
    if cover_c.get("show_report_version",True):
        mt_rows.append(("Version", meta.get("version","1.0")))

    tbl = doc.add_table(rows=len(mt_rows), cols=2)
    tbl.style = "Table Grid"
    for i,(k,v) in enumerate(mt_rows):
        row = tbl.rows[i]
        kc  = row.cells[0]
        vc  = row.cells[1]
        kc.text = k
        vc.text = str(v)
        for cell in [kc,vc]:
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        kc.paragraphs[0].runs[0].bold = True
        kc.paragraphs[0].runs[0].font.color.rgb = MUTED
        _set_cell_bg(kc, LIGHT_GRAY)

    if cover_c.get("custom_footer_note"):
        doc.add_paragraph()
        fn = doc.add_paragraph()
        fn.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fnr = fn.add_run(cover_c["custom_footer_note"])
        fnr.italic = True
        fnr.font.size = Pt(9)
        fnr.font.color.rgb = MUTED

    _add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════
    # TABLE OF CONTENTS
    # ══════════════════════════════════════════════════════════════════
    if struct.get("show_toc",True):
        toc_para = doc.add_paragraph()
        tr3 = toc_para.add_run("Table of Contents")
        tr3.bold = True
        tr3.font.size = Pt(16)
        tr3.font.color.rgb = DARK
        toc_para.paragraph_format.space_before = Pt(4)
        toc_para.paragraph_format.space_after  = Pt(6)
        _add_rule(doc, accent_hex)

        # Static sections
        for num, title, bm in [
            ("1.", "Executive Summary",    "sec_exec_summary"),
            ("2.", "Vulnerability Summary","sec_vuln_summary"),
            ("3.", "Findings Overview",    "sec_findings_overview"),
            ("4.", "Detailed Findings",    "sec_detailed_findings"),
        ]:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            nr = p.add_run(f"{num}  ")
            nr.font.size = Pt(10)
            nr.font.color.rgb = MUTED
            _add_hyperlink_to_bookmark(p, title, bm, font_size=10)

        doc.add_paragraph()
        vh = doc.add_paragraph()
        vhr = vh.add_run("Vulnerabilities")
        vhr.bold = True
        vhr.font.size = Pt(11)
        vhr.font.color.rgb = DARK
        vh.paragraph_format.space_after = Pt(4)

        # Vulnerability TOC rows
        if findings:
            tc_tbl = doc.add_table(rows=len(findings)+1, cols=4)
            tc_tbl.style = "Table Grid"

            # Header
            hdrs = ["ID","Title","Severity","CVSS"]
            for ci,h in enumerate(hdrs):
                cell = tc_tbl.rows[0].cells[ci]
                cell.text = h
                _set_cell_bg(cell, RGBColor(*[int(hdr_bg_hex[i:i+2],16) for i in (0,2,4)]))
                p2 = cell.paragraphs[0]
                p2.runs[0].bold = True
                p2.runs[0].font.size = Pt(9)
                p2.runs[0].font.color.rgb = WHITE
                p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

            for i,f in enumerate(findings):
                sev   = f.get("severity","Low")
                vid   = f.get("vuln_id","")
                bm_nm = f"vuln_{vid.replace('-','_')}"
                row   = tc_tbl.rows[i+1]

                # ID cell
                row.cells[0].text = vid
                row.cells[0].paragraphs[0].runs[0].font.size  = Pt(9)
                row.cells[0].paragraphs[0].runs[0].bold = True
                row.cells[0].paragraphs[0].runs[0].font.color.rgb = MUTED

                # Title cell — clickable
                tc_tbl.rows[i+1].cells[1].text = ""
                tp3 = tc_tbl.rows[i+1].cells[1].paragraphs[0]
                _add_hyperlink_to_bookmark(tp3, f.get("title",""), bm_nm, font_size=10)

                # Severity cell
                sc = row.cells[2]
                sc.text = sev
                sc.paragraphs[0].runs[0].bold = True
                sc.paragraphs[0].runs[0].font.size = Pt(9)
                sc.paragraphs[0].runs[0].font.color.rgb = _sev_rgb(sev)
                sc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                _set_cell_bg(sc, _sev_bg(sev))

                # CVSS cell
                row.cells[3].text = str(f.get("cvss_score",""))
                row.cells[3].paragraphs[0].runs[0].bold = True
                row.cells[3].paragraphs[0].runs[0].font.size = Pt(9)
                row.cells[3].paragraphs[0].runs[0].font.color.rgb = _sev_rgb(sev)
                row.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Alternating row bg
                if i % 2 == 0:
                    _set_cell_bg(row.cells[0], LIGHT_GRAY)
                    _set_cell_bg(row.cells[1], LIGHT_GRAY)
                    _set_cell_bg(row.cells[3], LIGHT_GRAY)

        _add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════
    # EXECUTIVE SUMMARY
    # ══════════════════════════════════════════════════════════════════
    if struct.get("show_executive_summary",True):
        es_p = doc.add_paragraph()
        _add_bookmark(es_p, next_bm(), "sec_exec_summary")
        esr = es_p.add_run("Executive Summary")
        esr.bold = True; esr.font.size = Pt(16); esr.font.color.rgb = DARK
        es_p.paragraph_format.space_after = Pt(6)
        _add_rule(doc, accent_hex)
        for line in exec_sum.split("\n"):
            if line.strip():
                bp = doc.add_paragraph(line.strip())
                bp.paragraph_format.space_after = Pt(4)
        doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════
    # VULNERABILITY SUMMARY CARDS
    # ══════════════════════════════════════════════════════════════════
    if struct.get("show_summary_cards",True):
        vs_p = doc.add_paragraph()
        _add_bookmark(vs_p, next_bm(), "sec_vuln_summary")
        vsr = vs_p.add_run("Vulnerability Summary")
        vsr.bold = True; vsr.font.size = Pt(16); vsr.font.color.rgb = DARK
        vs_p.paragraph_format.space_after = Pt(6)
        _add_rule(doc, accent_hex)

        counts  = _sev_counts(findings)
        overall = _overall_risk(findings)

        # Scorecard table
        sc_t = doc.add_table(rows=2, cols=5)
        sc_t.style = "Table Grid"
        sevs = ["Critical","High","Medium","Low","Informational"]
        for i,sev in enumerate(sevs):
            lc = sc_t.rows[0].cells[i]
            vc = sc_t.rows[1].cells[i]
            lc.text = sev
            vc.text = str(counts.get(sev,0))
            for cell in [lc,vc]:
                _set_cell_bg(cell, _sev_bg(sev))
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.paragraphs[0].runs[0].bold = True
                cell.paragraphs[0].runs[0].font.color.rgb = _sev_rgb(sev)
            vc.paragraphs[0].runs[0].font.size = Pt(20)

        doc.add_paragraph()
        rp = doc.add_paragraph()
        rp.add_run("Overall Risk Rating: ").font.size = Pt(11)
        rr = rp.add_run(overall)
        rr.bold = True; rr.font.size = Pt(11)
        rr.font.color.rgb = _sev_rgb(overall)
        rp.add_run(f"   |   Total Findings: {len(findings)}").font.size = Pt(11)
        doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════
    # FINDINGS OVERVIEW TABLE
    # ══════════════════════════════════════════════════════════════════
    if struct.get("show_findings_table",True):
        fo_p = doc.add_paragraph()
        _add_bookmark(fo_p, next_bm(), "sec_findings_overview")
        for_ = fo_p.add_run("Findings Overview")
        for_.bold = True; for_.font.size = Pt(16); for_.font.color.rgb = DARK
        fo_p.paragraph_format.space_after = Pt(6)
        _add_rule(doc, accent_hex)

        if findings:
            ov_t = doc.add_table(rows=len(findings)+1, cols=5)
            ov_t.style = "Table Grid"
            for ci,h in enumerate(["ID","Title","Severity","CVSS Score","Affected"]):
                cell = ov_t.rows[0].cells[ci]
                cell.text = h
                _set_cell_bg(cell, RGBColor(*[int(hdr_bg_hex[i:i+2],16) for i in (0,2,4)]))
                cell.paragraphs[0].runs[0].bold = True
                cell.paragraphs[0].runs[0].font.size = Pt(9)
                cell.paragraphs[0].runs[0].font.color.rgb = WHITE
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            for i,f in enumerate(findings):
                sev  = f.get("severity","Low")
                vals = [
                    f.get("vuln_id",""),
                    f.get("title",""),
                    sev,
                    str(f.get("cvss_score","")),
                    str(f.get("affected",""))[:60],
                ]
                for j,v in enumerate(vals):
                    cell = ov_t.rows[i+1].cells[j]
                    cell.text = v
                    cell.paragraphs[0].runs[0].font.size = Pt(9)
                    if j == 2:
                        cell.paragraphs[0].runs[0].bold = True
                        cell.paragraphs[0].runs[0].font.color.rgb = _sev_rgb(sev)
                        _set_cell_bg(cell, _sev_bg(sev))
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif j == 3:
                        cell.paragraphs[0].runs[0].bold = True
                        cell.paragraphs[0].runs[0].font.color.rgb = _sev_rgb(sev)
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif i % 2 == 0:
                        _set_cell_bg(cell, LIGHT_GRAY)

        _add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════
    # DETAILED FINDINGS — each on its own page
    # ══════════════════════════════════════════════════════════════════
    if struct.get("show_detailed_findings",True):
        df_p = doc.add_paragraph()
        _add_bookmark(df_p, next_bm(), "sec_detailed_findings")
        dfr = df_p.add_run("Detailed Findings")
        dfr.bold = True; dfr.font.size = Pt(16); dfr.font.color.rgb = DARK
        df_p.paragraph_format.space_after = Pt(6)
        _add_rule(doc, accent_hex)

        new_page = struct.get("each_vuln_new_page",True)
        for i,f in enumerate(findings):
            if i > 0 and new_page:
                _add_page_break(doc)
            _build_finding_section(doc, f, accent_hex, next_bm, DARK, WHITE,
                                   LIGHT_GRAY, MUTED, BODY_TEXT)

    # ══════════════════════════════════════════════════════════════════
    # DISCLAIMER
    # ══════════════════════════════════════════════════════════════════
    if struct.get("show_disclaimer",True):
        _add_page_break(doc)
        disc_p = doc.add_paragraph()
        discr  = disc_p.add_run("Disclaimer & Notes")
        discr.bold = True; discr.font.size = Pt(16); discr.font.color.rgb = DARK
        disc_p.paragraph_format.space_after = Pt(6)
        _add_rule(doc, accent_hex)

        disc = cfg.get("disclaimer",{})
        dp   = doc.add_paragraph(disc.get("text","This report is confidential."))
        dp.paragraph_format.space_after = Pt(6)

        if company.get("name"):
            cp2 = doc.add_paragraph()
            cr3 = cp2.add_run(f"Prepared by: {company['name']}")
            cr3.font.size = Pt(9); cr3.font.color.rgb = MUTED
            if company.get("website"):
                cp2.add_run(f"  |  {company['website']}").font.size = Pt(9)
            if company.get("email"):
                cp2.add_run(f"  |  {company['email']}").font.size = Pt(9)

        if disc.get("poc_note"):
            doc.add_paragraph()
            poc = doc.add_paragraph()
            pocr = poc.add_run(f"📸  {disc['poc_note']}")
            pocr.bold = True; pocr.font.size = Pt(9)
            pocr.font.color.rgb = RGBColor(210,153,34)

    os.makedirs(os.path.dirname(output_path) if os.path.dirname(output_path) else ".", exist_ok=True)
    doc.save(output_path)
    return output_path


# ── Single finding section ────────────────────────────────────────────────────

def _build_finding_section(doc, f, accent_hex, next_bm_fn,
                           DARK, WHITE, LIGHT_GRAY, MUTED, BODY_TEXT):
    sev   = f.get("severity","Low")
    vid   = f.get("vuln_id","")
    bm_nm = f"vuln_{vid.replace('-','_')}"

    # Title paragraph with bookmark
    tp = doc.add_paragraph()
    _add_bookmark(tp, next_bm_fn(), bm_nm)
    tp.paragraph_format.space_before = Pt(4)
    tp.paragraph_format.space_after  = Pt(4)

    # Severity badge + title
    badge_r = tp.add_run(f" {sev} ")
    badge_r.bold = True; badge_r.font.size = Pt(9)
    badge_r.font.color.rgb = _sev_rgb(sev)

    title_r = tp.add_run(f"  {vid}  —  {f.get('title','')}")
    title_r.bold = True; title_r.font.size = Pt(14)
    title_r.font.color.rgb = DARK

    # CVSS info line
    cv = doc.add_paragraph()
    cvr = cv.add_run(f"CVSS Score: {f.get('cvss_score','N/A')}    |    Vector: {f.get('cvss_vector','N/A')}")
    cvr.font.size = Pt(9); cvr.font.color.rgb = MUTED

    # CVSS breakdown — all 8 metrics in 2 rows of 4
    breakdown = cvss_breakdown(f.get("cvss_vector",""))
    if breakdown:
        items = list(breakdown.items())
        for chunk in [items[:4], items[4:]]:
            if not chunk: continue
            bt = doc.add_table(rows=2, cols=len(chunk))
            bt.style = "Table Grid"
            for ci,(label,val) in enumerate(chunk):
                lc = bt.rows[0].cells[ci]
                vc = bt.rows[1].cells[ci]
                lc.text = label
                vc.text = val
                for cell in [lc,vc]:
                    _set_cell_bg(cell, LIGHT_GRAY)
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                lc.paragraphs[0].runs[0].font.size = Pt(8)
                lc.paragraphs[0].runs[0].font.color.rgb = MUTED
                lc.paragraphs[0].runs[0].bold = True
                vc.paragraphs[0].runs[0].font.size = Pt(9.5)
                vc.paragraphs[0].runs[0].bold = True
            doc.add_paragraph()

    def labeled(label, text, mono=False):
        if not text or str(text).strip() in ("","NA","N/A"): return
        lp = doc.add_paragraph()
        lp.paragraph_format.space_after = Pt(2)
        lr = lp.add_run(label)
        lr.bold = True; lr.font.size = Pt(10)
        lr.font.color.rgb = MUTED
        for line in str(text).strip().split("\n"):
            if not line.strip(): continue
            bp = doc.add_paragraph()
            br = bp.add_run(line)
            br.font.size = Pt(10.5)
            if mono:
                br.font.name = "Consolas"
                br.font.size = Pt(9.5)
                _para_shade(bp, LIGHT_GRAY)
            bp.paragraph_format.space_after = Pt(2)

    labeled("Affected Component:",  f.get("affected",""))
    labeled("Description:",         f.get("description",""))
    labeled("Steps to Reproduce:",  f.get("steps_to_reproduce",""), mono=True)
    # POC Images — under Steps to Reproduce
    import json as _json2, os as _os2
    poc_paths2 = []
    try:
        poc_paths2 = _json2.loads(f.get("poc_images","[]") or "[]")
    except Exception:
        poc_paths2 = []
    if poc_paths2:
        from docx.shared import Cm as _Cm2
        ph = doc.add_paragraph()
        phr = ph.add_run("📸  POC Screenshots / Evidence")
        phr.bold = True; phr.font.size = Pt(10)
        phr.font.color.rgb = MUTED
        ph.paragraph_format.space_after = Pt(4)
        for img_path in poc_paths2:
            if img_path and _os2.path.exists(img_path):
                try:
                    from PIL import Image as PILImage
                    pil = PILImage.open(img_path)
                    w_px = pil.size[0]
                    dpi  = (pil.info.get("dpi",(96,96))[0]) or 96
                    w_cm = min((w_px / dpi) * 2.54, 15.0)
                    doc.add_picture(img_path, width=_Cm2(w_cm))
                    cap = doc.add_paragraph()
                    cr  = cap.add_run(f"Fig: {_os2.path.basename(img_path)}")
                    cr.italic = True; cr.font.size = Pt(8)
                    cr.font.color.rgb = MUTED
                except Exception as _ex:
                    ep = doc.add_paragraph()
                    ep.add_run(f"[Image: {_os2.path.basename(img_path)} — {_ex}]").italic = True

    labeled("Impact:",              f.get("impact",""))
    labeled("Recommendation:",      f.get("recommendations",""))
    if f.get("reference") and f.get("reference") not in ("NA","N/A"):
        labeled("References:", f.get("reference",""))

    # POC images embedded under Steps to Reproduce above

    doc.add_paragraph()
    _add_rule(doc, accent_hex)
